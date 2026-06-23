"""
DOCX document loader using python-docx.
"""

from typing import Dict, Any, List
from pathlib import Path
import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from backend.ingestion.loaders.base import BaseDocumentLoader, DocumentLoadError


class DOCXLoader(BaseDocumentLoader):
    """
    Loader for DOCX documents.
    
    Extracts text content, tables, and metadata from DOCX files using python-docx.
    """
    
    def __init__(self, include_tables: bool = True, include_headers: bool = True):
        """
        Initialize DOCX loader.
        
        Args:
            include_tables: Whether to extract table content
            include_headers: Whether to extract headers/footers
        """
        super().__init__()
        self.include_tables = include_tables
        self.include_headers = include_headers
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if file extension is .docx"""
        return file_extension.lower() in ['.docx']
    
    async def load(self, file_path: Path) -> Dict[str, Any]:
        """
        Load DOCX document and extract content.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with content, metadata, and structure
            
        Raises:
            DocumentLoadError: If DOCX loading fails
        """
        self._validate_file(file_path)
        
        try:
            # Load document
            doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = self._extract_docx_metadata(doc, file_path)
            
            # Extract content
            paragraphs = self._extract_paragraphs(doc)
            tables = self._extract_tables(doc) if self.include_tables else []
            headers_footers = self._extract_headers_footers(doc) if self.include_headers else []
            
            # Combine all content
            content_parts = []
            
            # Add paragraphs
            content_parts.extend([p["text"] for p in paragraphs if p["text"]])
            
            # Add tables
            if tables:
                content_parts.extend([t["text"] for t in tables if t["text"]])
            
            # Add headers/footers
            if headers_footers:
                content_parts.extend(headers_footers)
            
            content = "\n\n".join(content_parts)
            content = self._clean_text(content)
            
            self.logger.info(
                f"Successfully loaded DOCX: {file_path.name} "
                f"({len(paragraphs)} paragraphs, {len(tables)} tables, {len(content)} chars)"
            )
            
            return {
                "content": content,
                "metadata": metadata,
                "paragraphs": paragraphs,
                "tables": tables,
                "headers_footers": headers_footers
            }
            
        except Exception as e:
            raise DocumentLoadError(f"Failed to load DOCX: {e}")
    
    def _extract_docx_metadata(
        self, 
        doc: Document, 
        file_path: Path
    ) -> Dict[str, Any]:
        """
        Extract metadata from DOCX.
        
        Args:
            doc: python-docx Document object
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with DOCX metadata
        """
        # Get base metadata
        metadata = self._extract_base_metadata(file_path)
        
        # Add DOCX-specific metadata
        metadata.update({
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        })
        
        # Extract core properties if available
        if doc.core_properties:
            props = doc.core_properties
            
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.keywords:
                metadata["keywords"] = props.keywords
            if props.comments:
                metadata["comments"] = props.comments
            if props.created:
                metadata["created_date"] = props.created.isoformat()
            if props.modified:
                metadata["modified_date"] = props.modified.isoformat()
            if props.last_modified_by:
                metadata["last_modified_by"] = props.last_modified_by
        
        return metadata
    
    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract paragraphs from document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of paragraph dictionaries
        """
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            if text:  # Only include non-empty paragraphs
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else None,
                    "char_count": len(text)
                })
        
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract tables from document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of table dictionaries
        """
        tables = []
        
        for i, table in enumerate(doc.tables):
            try:
                # Extract table as text
                table_text = self._table_to_text(table)
                
                tables.append({
                    "index": i,
                    "text": table_text,
                    "row_count": len(table.rows),
                    "col_count": len(table.columns) if table.rows else 0,
                    "char_count": len(table_text)
                })
            except Exception as e:
                self.logger.warning(f"Failed to extract table {i}: {e}")
        
        return tables
    
    def _table_to_text(self, table: Table) -> str:
        """
        Convert table to text representation.
        
        Args:
            table: python-docx Table object
            
        Returns:
            Text representation of table
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        
        return "\n".join(rows)
    
    def _extract_headers_footers(self, doc: Document) -> List[str]:
        """
        Extract headers and footers from document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of header/footer texts
        """
        headers_footers = []
        
        try:
            # Extract from all sections
            for section in doc.sections:
                # Headers
                if section.header:
                    for para in section.header.paragraphs:
                        text = para.text.strip()
                        if text:
                            headers_footers.append(text)
                
                # Footers
                if section.footer:
                    for para in section.footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            headers_footers.append(text)
        except Exception as e:
            self.logger.warning(f"Failed to extract headers/footers: {e}")
        
        return headers_footers


# Made with Bob